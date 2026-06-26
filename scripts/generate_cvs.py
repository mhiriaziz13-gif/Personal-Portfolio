"""Generate the four downloadable CV variants from verified information only.

Run from the repository root:
    python3 scripts/generate_cvs.py

Requirements:
    python3 -m pip install python-docx
    LibreOffice (for PDF conversion)
"""
from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "public" / "cv"

INK = "0D1B2F"
MUTED = "58687B"
AMBER = "B85E25"
LINE = "D9E3EA"
SOFT = "F3F7FA"

CONTACT = "Sousse, Tunisia | mhiriaziz13@gmail.com | linkedin.com/in/ahmed-aziz-mhiri"
ENGLISH_HEADLINE = "Data-Driven Marketing & Commercial Analytics | Business Intelligence, Automation & Digital Growth"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs) -> None:
    """Set cell borders with kwargs such as bottom={val:'single', sz:'8', color:'D9E3EA'}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell_margin(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def set_font(run, size: float = 9.0, bold: bool = False, color: str = INK, name: str = "Arial") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure(doc: Document, *, letter: bool = False, ats: bool = False) -> None:
    sec = doc.sections[0]
    if letter:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
    else:
        sec.page_width = Inches(8.27)
        sec.page_height = Inches(11.69)
    sec.top_margin = Inches(0.45 if not ats else 0.42)
    sec.bottom_margin = Inches(0.42 if not ats else 0.42)
    sec.left_margin = Inches(0.52 if not ats else 0.58)
    sec.right_margin = Inches(0.52 if not ats else 0.58)

    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    normal.font.size = Pt(8.7 if not ats else 9.0)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.02

    for name in ['Title', 'Heading 1', 'Heading 2']:
        style = doc.styles[name]
        style.font.name = 'Arial'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')


def heading(doc: Document, text: str, *, compact: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if not compact else 5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '9')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), AMBER)
    pBdr.append(bottom)
    pPr.append(pBdr)
    r = p.add_run(text.upper())
    set_font(r, 9.5, True, INK)
    r.font.all_caps = True
    r.font.letter_spacing = Pt(0.6)


def add_header(doc: Document, headline: str, availability: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('AHMED AZIZ MHIRI')
    set_font(r, 21.5, True, INK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(headline)
    set_font(r2, 9.5, True, AMBER)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(1)
    r3 = p3.add_run(CONTACT)
    set_font(r3, 8.3, False, MUTED)
    if availability:
        p4 = doc.add_paragraph()
        p4.paragraph_format.space_after = Pt(3)
        r4 = p4.add_run(availability)
        set_font(r4, 8.2, True, MUTED)


def add_profile(doc: Document, text: str) -> None:
    heading(doc, 'Professional Profile')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, 8.7, False, MUTED)


def role_heading(doc: Document, title: str, company: str, date: str, location: str, *, font_size=9.2) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_font(r, font_size, True, INK)
    r2 = p.add_run(f' | {company}')
    set_font(r2, font_size, True, INK)
    r3 = p.add_run(f' | {date} | {location}')
    set_font(r3, 7.8, False, MUTED)


def bullet(doc: Document, text: str, *, font_size=8.45) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run('- ')
    set_font(r, font_size, True, AMBER)
    r2 = p.add_run(text)
    set_font(r2, font_size, False, INK)


def simple_line(doc: Document, title: str, text: str, *, font_size=8.45) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f'{title}: ')
    set_font(r, font_size, True, INK)
    r2 = p.add_run(text)
    set_font(r2, font_size, False, INK)


def skills_table(doc: Document, entries: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.52)
    table.columns[1].width = Inches(5.65)
    table.allow_autofit = False
    for i, (label, content) in enumerate(entries):
        cells = table.add_row().cells
        cells[0].width = Inches(1.52)
        cells[1].width = Inches(5.65)
        for cell in cells:
            set_cell_margin(cell, top=42, start=80, bottom=42, end=80)
            set_cell_border(cell, bottom={"val": "single", "sz": "3", "color": LINE})
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cells[0], SOFT)
        p1 = cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        set_font(r1, 7.75, True, INK)
        p2 = cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(content)
        set_font(r2, 8.0, False, INK)


def project_line(doc: Document, title: str, text: str, *, font_size=8.35) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.04)
    r = p.add_run(f'{title} - ')
    set_font(r, font_size, True, INK)
    r2 = p.add_run(text)
    set_font(r2, font_size, False, INK)


def education_table(doc: Document, entries: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for title, detail in entries:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margin(cell, top=45, start=70, bottom=45, end=70)
            set_cell_border(cell, bottom={"val": "single", "sz": "3", "color": LINE})
        p1 = cells[0].paragraphs[0]; p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(title); set_font(r1, 8.25, True, INK)
        p2 = cells[1].paragraphs[0]; p2.paragraph_format.space_after = Pt(0); p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(detail); set_font(r2, 7.75, False, MUTED)


def make_english() -> Document:
    doc = Document(); configure(doc)
    add_header(doc, ENGLISH_HEADLINE, 'Available for Europe-based opportunities from Summer 2027')
    add_profile(doc, 'Master’s student in Big Data Analytics & E-Commerce at IHEC Carthage with a Business Intelligence background. Interested in connecting operational, customer and marketing data to commercial decisions, digital customer journeys and reliable business processes. Experience spans process automation, digital marketing, hotel performance analysis and AI-enabled applications.')
    heading(doc, 'Core Expertise')
    skills_table(doc, [
        ('Data & BI', 'Data Analysis | Marketing Analytics | Commercial Analytics | Business Intelligence | KPI Analysis | Data Visualization | Financial Reporting | Excel'),
        ('Growth', 'Digital Marketing | Customer Insights | Customer Journey | Local SEO | Email Marketing | Paid Social | Social Media Strategy | E-Commerce'),
        ('Automation', 'UiPath | Process Automation | Business Rules Automation | JSON | HTML Reporting | Angular | Spring Boot | REST APIs | RAG'),
        ('Strengths', 'Commercial Thinking | Problem Solving | Cross-Functional Collaboration | Process Improvement | Structured Communication'),
    ])
    heading(doc, 'Professional Experience')
    role_heading(doc, 'Head of IT Services - Process Automation & Business Systems', 'Sunshine Vacances France', 'Jul 2025 - Present', 'Sousse, Tunisia')
    for x in [
        'Designed and expanded a UiPath workflow for invoice validation and reconciliation across invoices, vouchers, reservations and stay-related data.',
        'Automated commercial rules covering room rates, board types, discounts, supplements and special offers.',
        'Generated structured JSON outputs and HTML reports to support audit, review and operational follow-up; adapted controls to multiple agency and hotel configurations.',
    ]: bullet(doc, x)
    role_heading(doc, 'Digital Marketing & Automation Consultant', 'Confidential Client - Men’s Barbershop', 'Feb 2025 - Jul 2025', 'Noisy-le-Grand, France')
    for x in [
        'Built and managed a tailored website with online booking and activity monitoring to improve the booking journey.',
        'Supported local SEO, customer communication, social content, paid social activity, online reviews and the wider digital presence.',
        'Facilitated a partnership with Planity to streamline booking operations and improve customer touchpoints.',
    ]: bullet(doc, x)
    role_heading(doc, 'AI & Full-Stack Development Intern', 'VERMEG for Banking & Insurance Software', 'Feb 2025 - May 2025', 'Tunis, Tunisia')
    for x in [
        'Developed key e-learning platform features including course enrolment, progress tracking, dashboards, internal event booking and real-time notifications.',
        'Integrated a locally deployed LLaMA 3.2 assistant through Ollama with RAG retrieval for PDF and CSV knowledge sources.',
        'Contributed to secure, scalable microservices features and measures addressing prompt injection, malicious files and unsafe links.',
    ]: bullet(doc, x)
    role_heading(doc, 'Commercial & Digital Marketing Manager', 'Maison Salina', 'Apr 2025 - Sep 2025', 'Sousse, Tunisia')
    for x in [
        'Developed digital marketing initiatives aligned with commercial objectives, brand positioning and customer-facing communication.',
        'Supported strategic collaborations, online visibility and the monitoring of commercial priorities and digital actions.',
    ]: bullet(doc, x)
    role_heading(doc, 'Management Controller', 'El Mouradi Hotels - El Mouradi Kantaoui Club', 'Jul 2024 - Sep 2024', 'Sousse, Tunisia')
    for x in ['Analysed occupancy, operational costs and revenue-related KPIs; contributed to budget preparation, variance analysis and management reporting.', 'Supported identification of cost-control opportunities across procurement, inventory and operational processes.']: bullet(doc, x)
    role_heading(doc, 'Full-Stack Development Intern', 'Arabsoft', 'Jun 2024 - Aug 2024', 'Tunis, Tunisia')
    bullet(doc, 'Designed a responsive Angular interface and Spring Boot REST APIs for a books, users and borrowing-management application.')
    role_heading(doc, 'Management Control Intern', 'El Mouradi Hotels - El Mouradi Palace', 'Jun 2023 - Sep 2023', 'Sousse, Tunisia')
    bullet(doc, 'Analysed expenses, operating indicators and budget variances; contributed to management reporting, cost-control initiatives and KPI monitoring.')
    heading(doc, 'Selected Projects', compact=True)
    project_line(doc, 'RPA for Invoice Control & Booking Reconciliation', 'UiPath-based validation workflow connecting invoice data, bookings, vouchers, rates, discounts and supplements to structured JSON and HTML audit outputs.')
    project_line(doc, 'Data-Driven Digital Transformation for a Men’s Barbershop', 'Website, online booking, activity monitoring, local SEO, social initiatives and more structured customer touchpoints.')
    project_line(doc, 'AI-Ready E-Learning Platform', 'Contributed to Angular and Spring Cloud features, RAG-enabled knowledge retrieval via Ollama and LLaMA 3.2, security controls and productivity tools.')
    heading(doc, 'Education & Certification', compact=True)
    education_table(doc, [
        ('Master’s Degree - Big Data Analytics & E-Commerce, IHEC Carthage', 'Oct 2025 - Jun 2027 | Expected Jun 2027'),
        ('Bachelor’s Degree - Business Intelligence, IHEC Carthage', 'Jan 2021 - Jun 2025'),
        ('Fundamentals of Digital Marketing - Google', 'Certification'),
    ])
    simple_line(doc, 'Languages', 'French | English', font_size=8.2)
    return doc


def make_french() -> Document:
    doc = Document(); configure(doc)
    add_header(doc, 'Analytique Marketing, Analyse Commerciale, Business Intelligence, Automatisation des Processus & Croissance Digitale', 'Disponible pour des opportunités en Europe à partir de l’été 2027')
    heading(doc, 'Profil Professionnel')
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Étudiant en Master Big Data Analytics & E-Commerce à l’IHEC Carthage, titulaire d’une licence en Business Intelligence. Intéressé par l’exploitation des données opérationnelles, clients et marketing pour soutenir les décisions commerciales, améliorer le parcours client et fiabiliser les processus. Expériences en automatisation, transformation digitale, analyse de performance hôtelière et applications enrichies par l’IA.")
    set_font(r, 8.7, False, MUTED)
    heading(doc, 'Compétences Clés')
    skills_table(doc, [
        ('Data & BI', 'Analyse de données | Analytique Marketing | Analyse Commerciale | Business Intelligence | Analyse des KPI | Visualisation | Reporting financier | Excel'),
        ('Croissance', 'Marketing digital | Connaissance client | Parcours client | SEO local | Email marketing | Social Ads | Stratégie réseaux sociaux | E-Commerce'),
        ('Technologies', 'UiPath | Automatisation des Processus | Automatisation des règles métier | JSON | Reporting HTML | Angular | Spring Boot | API REST | RAG'),
        ('Atouts', 'Vision commerciale | Résolution de problèmes | Collaboration transverse | Amélioration des processus | Communication structurée'),
    ])
    heading(doc, 'Expériences Professionnelles')
    role_heading(doc, 'Responsable des Services IT - Automatisation des Processus & Systèmes Métier', 'Sunshine Vacances France', 'Juil. 2025 - Présent', 'Sousse, Tunisie')
    for x in [
        'Conception et extension d’un workflow UiPath de contrôle et de rapprochement des factures, vouchers, réservations et données de séjour.',
        'Automatisation de règles commerciales liées aux tarifs, pensions, remises, suppléments et offres spéciales.',
        'Production de sorties JSON structurées et de rapports HTML pour l’audit, le suivi opérationnel et la traçabilité; adaptation à plusieurs processus d’agence et configurations hôtelières.',
    ]: bullet(doc, x)
    role_heading(doc, 'Consultant Marketing Digital & Automatisation', 'Client confidentiel - Barbershop masculin', 'Fév. 2025 - Juil. 2025', 'Noisy-le-Grand, France')
    for x in [
        'Création et gestion d’un site web avec réservation en ligne et suivi de l’activité pour fluidifier le parcours de prise de rendez-vous.',
        'Soutien au SEO local, à la communication client, aux contenus sociaux, aux campagnes payantes, aux avis en ligne et à la visibilité digitale.',
        'Mise en relation avec Planity afin de simplifier les opérations de réservation et structurer les points de contact client.',
    ]: bullet(doc, x)
    role_heading(doc, 'Stagiaire IA & Développement Full-Stack', 'VERMEG for Banking & Insurance Software', 'Fév. 2025 - Mai 2025', 'Tunis, Tunisie')
    for x in [
        'Développement de fonctionnalités clés d’une plateforme e-learning: inscription aux cours, suivi de progression, tableaux de bord, réservation d’événements internes et notifications temps réel.',
        'Intégration d’un assistant LLaMA 3.2 déployé localement via Ollama, avec RAG pour l’interrogation de connaissances PDF et CSV.',
        'Contribution à l’architecture microservices sécurisée et évolutive, notamment sur les risques liés au prompt injection, aux fichiers malveillants et aux liens non sûrs.',
    ]: bullet(doc, x)
    role_heading(doc, 'Responsable Commercial & Marketing Digital', 'Maison Salina', 'Avr. 2025 - Sept. 2025', 'Sousse, Tunisie')
    for x in [
        'Déploiement d’initiatives de marketing digital cohérentes avec les objectifs commerciaux, le positionnement et la communication de marque.',
        'Contribution aux collaborations stratégiques, à la visibilité en ligne et au suivi des priorités commerciales et digitales.',
    ]: bullet(doc, x)
    role_heading(doc, 'Contrôleur de Gestion', 'El Mouradi Hotels - El Mouradi Kantaoui Club', 'Juil. 2024 - Sept. 2024', 'Sousse, Tunisie')
    for x in ['Analyse du taux d’occupation, des coûts opérationnels et des KPI de revenus; contribution au budget, aux analyses d’écarts et au reporting de gestion.', 'Appui à l’identification de leviers de maîtrise des coûts en achats, stocks et opérations.']: bullet(doc, x)
    role_heading(doc, 'Stagiaire Développement Full-Stack', 'Arabsoft', 'Juin 2024 - Août 2024', 'Tunis, Tunisie')
    bullet(doc, 'Conception d’une interface Angular responsive et d’API REST Spring Boot pour une application de gestion des livres, utilisateurs et emprunts.')
    role_heading(doc, 'Stagiaire Contrôle de Gestion', 'El Mouradi Hotels - El Mouradi Palace', 'Juin 2023 - Sept. 2023', 'Sousse, Tunisie')
    bullet(doc, 'Analyse des dépenses, indicateurs opérationnels et écarts budgétaires; contribution au reporting, au contrôle des coûts et au suivi des KPI.')
    heading(doc, 'Projets Sélectionnés', compact=True)
    project_line(doc, 'RPA - Contrôle des factures et rapprochement des réservations', 'Workflow UiPath reliant factures, réservations, vouchers, tarifs, remises et suppléments à des sorties JSON et HTML pour l’audit.')
    project_line(doc, 'Transformation digitale d’un barbershop masculin', 'Site web, prise de rendez-vous en ligne, suivi d’activité, SEO local, initiatives sociales et parcours client plus structuré.')
    project_line(doc, 'Plateforme e-learning prête pour l’IA', 'Contribution aux fonctionnalités Angular et Spring Cloud, au RAG via Ollama et LLaMA 3.2, aux contrôles de sécurité et aux outils de productivité.')
    heading(doc, 'Formation & Certification', compact=True)
    education_table(doc, [
        ('Master Big Data Analytics & E-Commerce - IHEC Carthage', 'Oct. 2025 - Juin 2027 | Diplôme prévu: Juin 2027'),
        ('Licence Business Intelligence - IHEC Carthage', 'Jan. 2021 - Juin 2025'),
        ('Fundamentals of Digital Marketing - Google', 'Certification'),
    ])
    simple_line(doc, 'Langues', 'Français | Anglais', font_size=8.2)
    return doc


def make_ats() -> Document:
    doc = Document(); configure(doc, ats=True)
    add_header(doc, 'Marketing Analyst | Commercial Analyst | Business Intelligence Analyst | Digital Marketing Analyst | Revenue Operations Analyst | Process Automation Analyst | Junior Business Analyst', 'Available for Europe-based opportunities from Summer 2027')
    heading(doc, 'Professional Summary')
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('Master’s student in Big Data Analytics & E-Commerce with a Business Intelligence background. Brings experience across marketing analytics, commercial analytics, KPI analysis, digital marketing, local SEO, customer journey improvement, process automation, financial reporting and AI-enabled application development. Interested in roles that connect data, operations, customer signals and commercial performance.')
    set_font(r, 8.9, False, INK)
    heading(doc, 'Core Skills')
    for label, content in [
        ('Data & Analytics', 'Data Analysis; Marketing Analytics; Commercial Analytics; Business Intelligence; KPI Analysis; Data Visualization; Financial Reporting; Excel'),
        ('Marketing & Growth', 'Digital Marketing; Customer Insights; Customer Journey; Local SEO; Email Marketing; Paid Social; Social Media Strategy; E-Commerce'),
        ('Automation & Technology', 'UiPath; Process Automation; Business Rules Automation; JSON; HTML Reporting; Angular; Spring Boot; REST APIs; RAG; Ollama; LLaMA 3.2'),
        ('Business Strengths', 'Commercial Thinking; Problem Solving; Cross-Functional Collaboration; Process Improvement; Structured Communication'),
    ]: simple_line(doc, label, content, font_size=8.45)
    heading(doc, 'Professional Experience')
    role_heading(doc, 'Head of IT Services - Process Automation & Business Systems', 'Sunshine Vacances France', 'July 2025 - Present', 'Sousse, Tunisia', font_size=9.0)
    for x in ['Designed and expanded RPA workflow for invoice validation and booking reconciliation using invoices, vouchers, reservations and stay-related data.', 'Automated business rules covering room rates, board types, discounts, supplements and special offers.', 'Generated JSON outputs and HTML reports for audit, review and operational follow-up; improved process traceability and reliability.']: bullet(doc, x, font_size=8.45)
    role_heading(doc, 'Commercial & Digital Marketing Manager', 'Maison Salina', 'April 2025 - September 2025', 'Sousse, Tunisia', font_size=9.0)
    for x in ['Developed digital marketing initiatives aligned with commercial objectives and brand positioning.', 'Supported digital presence, customer-facing communication, strategic collaborations and monitoring of commercial priorities.']: bullet(doc, x, font_size=8.45)
    role_heading(doc, 'Digital Marketing & Automation Consultant', 'Confidential Client - Men’s Barbershop', 'February 2025 - July 2025', 'Noisy-le-Grand, France', font_size=9.0)
    for x in ['Built and managed a website with online booking and activity monitoring.', 'Supported local SEO, social content, paid social, reviews, email marketing and customer communication.', 'Facilitated Planity partnership to streamline booking operations and improve customer journey.']: bullet(doc, x, font_size=8.45)
    role_heading(doc, 'AI & Full-Stack Development Intern', 'VERMEG for Banking & Insurance Software', 'February 2025 - May 2025', 'Tunis, Tunisia', font_size=9.0)
    for x in ['Developed e-learning platform features: enrolment, progress tracking, dashboards, event booking and real-time notifications.', 'Integrated locally deployed LLaMA 3.2 assistant through Ollama with RAG retrieval for PDF and CSV knowledge.', 'Contributed to microservices, secure user management, monitoring and AI security controls.']: bullet(doc, x, font_size=8.45)
    role_heading(doc, 'Management Controller', 'El Mouradi Hotels - El Mouradi Kantaoui Club', 'July 2024 - September 2024', 'Sousse, Tunisia', font_size=9.0)
    bullet(doc, 'Analysed occupancy, operational cost and revenue KPIs; contributed to budgeting, variance analysis, management reporting and cost-control opportunities.', font_size=8.45)
    role_heading(doc, 'Full-Stack Development Intern', 'Arabsoft', 'June 2024 - August 2024', 'Tunis, Tunisia', font_size=9.0)
    bullet(doc, 'Designed Angular user interface and Spring Boot REST APIs for management of books, users and borrowing operations.', font_size=8.45)
    role_heading(doc, 'Management Control Intern', 'El Mouradi Hotels - El Mouradi Palace', 'June 2023 - September 2023', 'Sousse, Tunisia', font_size=9.0)
    bullet(doc, 'Analysed expenses, operational indicators and budget variances; supported management reporting, cost control and KPI monitoring.', font_size=8.45)
    heading(doc, 'Selected Projects', compact=True)
    project_line(doc, 'RPA for Invoice Control & Booking Reconciliation', 'UiPath workflow for invoice and booking validation, commercial-rule automation, structured JSON and HTML audit reporting.', font_size=8.4)
    project_line(doc, 'Data-Driven Digital Transformation for a Men’s Barbershop', 'Website, booking journey, local SEO, social media, customer communication and activity monitoring.', font_size=8.4)
    project_line(doc, 'AI-Ready E-Learning Platform', 'Angular, Spring Boot, Spring Cloud, Keycloak, Kafka, Ollama, RAG, monitoring and security tooling.', font_size=8.4)
    heading(doc, 'Education')
    simple_line(doc, 'Master’s Degree', 'Big Data Analytics & E-Commerce, IHEC Carthage | October 2025 - June 2027 | Expected graduation: June 2027', font_size=8.4)
    simple_line(doc, 'Bachelor’s Degree', 'Business Intelligence, IHEC Carthage | January 2021 - June 2025', font_size=8.4)
    heading(doc, 'Certification & Languages', compact=True)
    simple_line(doc, 'Certification', 'Fundamentals of Digital Marketing - Google', font_size=8.4)
    simple_line(doc, 'Languages', 'French | English', font_size=8.4)
    return doc


def make_canada() -> Document:
    doc = Document(); configure(doc, letter=True)
    add_header(doc, ENGLISH_HEADLINE, 'Open to relocation | Available from Summer 2027')
    add_profile(doc, 'Master’s student in Big Data Analytics & E-Commerce with a Business Intelligence background. Applies data, automation and digital marketing to commercial performance, customer journeys and operational decision-making. Offers transferable experience in RPA, KPI analysis, financial reporting, local digital growth, web platforms and RAG-enabled applications.')
    heading(doc, 'Transferable Skills')
    skills_table(doc, [
        ('Analytics', 'Data Analysis | Marketing Analytics | Commercial Analytics | Business Intelligence | KPI Analysis | Data Visualization | Financial Reporting | Excel'),
        ('Customer Growth', 'Digital Marketing | Customer Insights | Customer Journey | Local SEO | Email Marketing | Paid Social | Social Media Strategy | E-Commerce'),
        ('Automation & Tech', 'UiPath | Process Automation | Business Rules Automation | JSON | HTML Reporting | Angular | Spring Boot | REST APIs | RAG'),
        ('Strengths', 'Commercial Thinking | Problem Solving | Cross-Functional Collaboration | Process Improvement | Structured Communication'),
    ])
    heading(doc, 'Relevant Experience')
    role_heading(doc, 'Head of IT Services - Process Automation & Business Systems', 'Sunshine Vacances France', 'Jul 2025 - Present', 'Sousse, Tunisia')
    for x in ['Designed and expanded an RPA workflow for invoice validation and reconciliation across invoices, vouchers, reservations and stay-related data.', 'Automated commercial rules covering room rates, board types, discounts, supplements and special offers.', 'Created structured JSON outputs and HTML reports that support audit, review, operational follow-up and traceability.']: bullet(doc, x)
    role_heading(doc, 'Digital Marketing & Automation Consultant', 'Confidential Client - Men’s Barbershop', 'Feb 2025 - Jul 2025', 'Noisy-le-Grand, France')
    for x in ['Built and managed a tailored website with online booking and activity monitoring.', 'Supported local SEO, social media, email marketing, paid social, online reviews and customer communication.', 'Facilitated Planity partnership to improve booking operations and customer touchpoints.']: bullet(doc, x)
    role_heading(doc, 'Commercial & Digital Marketing Manager', 'Maison Salina', 'Apr 2025 - Sep 2025', 'Sousse, Tunisia')
    for x in ['Developed digital marketing initiatives aligned with commercial objectives, brand positioning and customer communication.', 'Supported strategic collaborations, online visibility and commercial-priority monitoring.']: bullet(doc, x)
    role_heading(doc, 'AI & Full-Stack Development Intern', 'VERMEG for Banking & Insurance Software', 'Feb 2025 - May 2025', 'Tunis, Tunisia')
    for x in ['Developed key learning-platform features: enrolment, progress tracking, dashboards, internal event booking and real-time notifications.', 'Integrated a locally deployed LLaMA 3.2 assistant via Ollama with RAG retrieval for PDF and CSV knowledge sources.', 'Contributed to secure, scalable microservices features and AI safety measures.']: bullet(doc, x)
    role_heading(doc, 'Management Controller', 'El Mouradi Hotels - El Mouradi Kantaoui Club', 'Jul 2024 - Sep 2024', 'Sousse, Tunisia')
    bullet(doc, 'Analysed occupancy, operational costs and revenue-related KPIs; supported annual budgeting, variance analysis, financial reporting and cost-control opportunities.')
    role_heading(doc, 'Full-Stack Development Intern', 'Arabsoft', 'Jun 2024 - Aug 2024', 'Tunis, Tunisia')
    bullet(doc, 'Designed responsive Angular user interface and Spring Boot REST APIs for a management application covering books, users and borrowing operations.')
    role_heading(doc, 'Management Control Intern', 'El Mouradi Hotels - El Mouradi Palace', 'Jun 2023 - Sep 2023', 'Sousse, Tunisia')
    bullet(doc, 'Analysed expenses, operational indicators and budget variances; contributed to reporting, cost-control initiatives and KPI follow-up.')
    heading(doc, 'Selected Project Value', compact=True)
    project_line(doc, 'RPA for Invoice Control & Booking Reconciliation', 'Automated business-rule validation and generated structured outputs for more repeatable operational control and audit review.')
    project_line(doc, 'Digital Transformation for a Local Service Business', 'Connected website, online booking, local discovery, social visibility and customer communication.')
    project_line(doc, 'AI-Ready E-Learning Platform', 'Contributed to a secure learning platform with RAG-enabled knowledge retrieval and productivity features.')
    heading(doc, 'Education & Certification', compact=True)
    education_table(doc, [
        ('Master’s Degree - Big Data Analytics & E-Commerce, IHEC Carthage', 'Oct 2025 - Jun 2027 | Expected Jun 2027'),
        ('Bachelor’s Degree - Business Intelligence, IHEC Carthage', 'Jan 2021 - Jun 2025'),
        ('Fundamentals of Digital Marketing - Google', 'Certification'),
    ])
    simple_line(doc, 'Languages', 'French | English', font_size=8.2)
    return doc


def convert_pdf(docx_path: Path) -> Path:
    with tempfile.TemporaryDirectory(prefix="lo-profile-") as temp:
        profile_uri = Path(temp).as_uri()
        env = os.environ.copy()
        env['HOME'] = temp
        cmd = ['libreoffice', f'-env:UserInstallation={profile_uri}', '--headless', '--convert-to', 'pdf', '--outdir', str(OUTPUT), str(docx_path)]
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, env=env)
        expected = OUTPUT / f'{docx_path.stem}.pdf'
        if result.returncode != 0 or not expected.exists():
            raise RuntimeError(f'LibreOffice failed for {docx_path.name}:\n{result.stdout}\n{result.stderr}')
    return expected


def create(name: str, document: Document) -> None:
    docx_path = OUTPUT / f'{name}.docx'
    document.core_properties.author = 'Ahmed Aziz Mhiri'
    document.core_properties.title = name
    document.core_properties.subject = 'Professional CV'
    document.core_properties.comments = 'Generated from verified portfolio information.'
    document.save(docx_path)
    convert_pdf(docx_path)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for child in OUTPUT.iterdir():
        if child.suffix.lower() in {'.pdf', '.docx'}:
            child.unlink()
    create('Ahmed_Aziz_Mhiri_CV_English', make_english())
    create('Ahmed_Aziz_Mhiri_CV_Français', make_french())
    create('Ahmed_Aziz_Mhiri_CV_ATS', make_ats())
    create('Ahmed_Aziz_Mhiri_CV_Canada', make_canada())
    print('Generated CV files:')
    for file in sorted(OUTPUT.iterdir()):
        print(f'- {file.name}: {file.stat().st_size // 1024} KB')

if __name__ == '__main__':
    main()
